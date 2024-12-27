import pytest
import asyncio
from pathlib import Path
from docx import Document
from PyPDF2 import PdfWriter, PdfReader
import io
from reportlab.pdfgen import canvas
from app.services.extractor import DocumentExtractor
from app.services.chunker import DocumentChunker
from app.services.embedding import EmbeddingService

def create_large_pdf():
    """100페이지 이상의 PDF 파일 생성"""
    output = PdfWriter()
    
    # 메모리에 PDF 페이지 생성
    for i in range(150):  # 150 페이지
        packet = io.BytesIO()
        can = canvas.Canvas(packet)
        
        # 다양한 언어로 텍스트 추가
        can.drawString(100, 800, f"Page {i+1}")
        can.drawString(100, 750, "English text for testing")
        can.drawString(100, 700, "한글 테스트 텍스트")
        can.drawString(100, 650, "日本語テストテキスト")
        can.drawString(100, 600, "中文测试文本")
        can.drawString(100, 550, "Texte de test en français")
        can.drawString(100, 500, "Texto de prueba en español")
        
        can.save()
        packet.seek(0)
        new_pdf = PdfReader(packet)
        output.add_page(new_pdf.pages[0])
    
    return output

def create_large_docx():
    """100페이지 이상의 DOCX 파일 생성"""
    doc = Document()
    
    for i in range(150):  # 150 페이지
        doc.add_heading(f'Page {i+1}', 0)
        doc.add_paragraph('English paragraph for testing')
        doc.add_paragraph('한글 테스트 문단')
        doc.add_paragraph('日本語テスト段落')
        doc.add_paragraph('中文测试段落')
        doc.add_paragraph('Paragraphe de test en français')
        doc.add_paragraph('Párrafo de prueba en español')
        doc.add_page_break()
    
    return doc

@pytest.fixture
async def services():
    """테스트용 서비스 생성"""
    extractor = DocumentExtractor()
    chunker = DocumentChunker(chunk_size=200, chunk_overlap=50)  
    embedding_service = EmbeddingService()
    await embedding_service.initialize()
    
    return extractor, chunker, embedding_service

@pytest.mark.asyncio
async def test_large_pdf_processing(services, tmp_path):
    """대용량 PDF 파일 처리 테스트"""
    extractor, chunker, embedding_service = await services
    
    # 대용량 PDF 생성
    pdf_path = tmp_path / "large_test.pdf"
    pdf = create_large_pdf()
    with open(pdf_path, "wb") as f:
        pdf.write(f)
    
    # 파일 읽기
    with open(pdf_path, "rb") as f:
        content = f.read()
    
    # 텍스트 추출
    start_time = asyncio.get_event_loop().time()
    text = extractor.extract_text(content, "application/pdf")
    extraction_time = asyncio.get_event_loop().time() - start_time
    
    assert text is not None
    assert len(text) > 1000  # 충분한 텍스트가 추출되었는지 확인
    
    # 청킹
    start_time = asyncio.get_event_loop().time()
    chunks = chunker.create_chunks(text)
    chunking_time = asyncio.get_event_loop().time() - start_time
    
    assert len(chunks) > 50  # 충분한 수의 청크가 생성되었는지 확인
    
    # 임베딩 생성 및 저장
    start_time = asyncio.get_event_loop().time()
    metadata = [{"page": i, "source": "large_test.pdf"} for i in range(len(chunks))]
    await embedding_service.upsert_texts(chunks[:10], metadata[:10])  # 처음 10개 청크만 테스트
    embedding_time = asyncio.get_event_loop().time() - start_time
    
    print(f"\nPerformance Metrics for Large PDF:")
    print(f"Extraction Time: {extraction_time:.2f}s")
    print(f"Chunking Time: {chunking_time:.2f}s")
    print(f"Embedding Time (10 chunks): {embedding_time:.2f}s")
    print(f"Total Chunks: {len(chunks)}")

@pytest.mark.asyncio
async def test_large_docx_processing(services, tmp_path):
    """대용량 DOCX 파일 처리 테스트"""
    extractor, chunker, embedding_service = await services
    
    # 대용량 DOCX 생성
    docx_path = tmp_path / "large_test.docx"
    doc = create_large_docx()
    doc.save(docx_path)
    
    # 파일 읽기
    with open(docx_path, "rb") as f:
        content = f.read()
    
    # 텍스트 추출
    start_time = asyncio.get_event_loop().time()
    text = extractor.extract_text(content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    extraction_time = asyncio.get_event_loop().time() - start_time
    
    assert text is not None
    assert len(text) > 1000
    
    # 청킹
    start_time = asyncio.get_event_loop().time()
    chunks = chunker.create_chunks(text)
    chunking_time = asyncio.get_event_loop().time() - start_time
    
    assert len(chunks) > 50
    
    # 임베딩 생성 및 저장
    start_time = asyncio.get_event_loop().time()
    metadata = [{"page": i, "source": "large_test.docx"} for i in range(len(chunks))]
    await embedding_service.upsert_texts(chunks[:10], metadata[:10])  # 처음 10개 청크만 테스트
    embedding_time = asyncio.get_event_loop().time() - start_time
    
    print(f"\nPerformance Metrics for Large DOCX:")
    print(f"Extraction Time: {extraction_time:.2f}s")
    print(f"Chunking Time: {chunking_time:.2f}s")
    print(f"Embedding Time (10 chunks): {embedding_time:.2f}s")
    print(f"Total Chunks: {len(chunks)}")

@pytest.mark.asyncio
async def test_multilingual_search(services, tmp_path):
    """다국어 문서 검색 테스트"""
    extractor, chunker, embedding_service = await services
    
    # 다국어 텍스트 생성
    multilingual_texts = [
        "This is an English test document",
        "이것은 한글 테스트 문서입니다",
        "これは日本語のテスト文書です",
        "这是中文测试文档",
        "Ceci est un document de test en français",
        "Este es un documento de prueba en español"
    ]
    
    # 임베딩 생성 및 저장
    metadata = [{"lang": lang, "id": str(i)} for i, lang in enumerate(["en", "ko", "ja", "zh", "fr", "es"])]
    await embedding_service.upsert_texts(multilingual_texts, metadata)
    
    # 각 언어로 검색 테스트
    queries = [
        ("english document", "en"),
        ("한글 문서", "ko"),
        ("日本語の文書", "ja"),
        ("中文文档", "zh"),
        ("document français", "fr"),
        ("documento español", "es")
    ]
    
    for query, expected_lang in queries:
        results = await embedding_service.query(query, top_k=1)
        assert results is not None
        assert len(results) == 1
        assert results[0].metadata["lang"] == expected_lang
