import pytest
import os
from pathlib import Path
from datetime import datetime
from docx import Document
from uuid import uuid4
from tempfile import SpooledTemporaryFile
from fastapi import UploadFile
from sqlalchemy.ext.asyncio import AsyncSession
from app.services.storage import StorageService
from app.services.document import DocumentService
from app.services.rag import RAGService

@pytest.fixture
def test_docx_file():
    # 테스트용 DOCX 파일 생성
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test document created for testing purposes.')
    doc.add_paragraph('It contains multiple paragraphs with different content.')
    doc.add_heading('Department Information', level=1)
    doc.add_paragraph('IT Department: Development team')
    doc.add_paragraph('HR Department: Human resources team')
    
    file_path = Path('test_data.docx')
    doc.save(file_path)
    yield file_path
    
    # 테스트 후 파일 삭제
    if file_path.exists():
        file_path.unlink()

class MockUploadFile(UploadFile):
    def __init__(self, filename: str):
        # 파일을 SpooledTemporaryFile에 복사
        spooled = SpooledTemporaryFile()
        with open(filename, 'rb') as f:
            spooled.write(f.read())
        spooled.seek(0)
        
        super().__init__(
            filename=os.path.basename(filename),
            file=spooled,
            headers={"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
        )

@pytest.mark.asyncio
async def test_docx_processing_pipeline(test_docx_file, async_session):
    session = await anext(async_session)
    try:
        # 1. Storage Service를 통한 파일 업로드
        storage_service = StorageService()
        with open(test_docx_file, 'rb') as f:
            file_content = f.read()
        
        # 파일 이름으로 destination_blob_name 생성
        current_time = datetime.now().strftime("%Y%m%d_%H%M%S")
        destination_blob_name = f"test_uploads/{current_time}_{test_docx_file.name}"
        
        storage_path = await storage_service.upload_file(
            file_content=file_content,
            destination_blob_name=destination_blob_name
        )
        assert storage_path, "파일 업로드 실패"
        print(f"1. 파일 업로드 성공: {storage_path}")
        
        # 2. Document Service를 통한 문서 처리
        document_service = DocumentService(session)
        mock_file = MockUploadFile(filename=str(test_docx_file))
        project_id = uuid4()
        document_ids = await document_service.upload_documents(project_id, [mock_file])
        assert document_ids and len(document_ids) > 0, "문서 생성 실패"
        document = await document_service.get_document(document_ids[0])
        assert document, "문서 조회 실패"
        print(f"2. 문서 생성 성공: {document.id}")
        
        # 3. RAG Service를 통한 문서 처리 및 임베딩
        rag_service = RAGService(session)
        chunks = await rag_service.process_document(document)
        assert chunks and len(chunks) > 0, "RAG 처리 실패"
        print(f"3. RAG 처리 성공: {len(chunks)} 개의 청크 생성 및 임베딩 완료")
        
        # 4. 간단한 검색 테스트
        results = await rag_service.query(
            "IT department",
            k=1
        )
        assert results and len(results) > 0, "검색 결과 없음"
        print(f"4. 검색 테스트 성공: {results}")
        
        print("\n전체 테스트 성공!")
        
    except Exception as e:
        pytest.fail(f"테스트 실패: {str(e)}")
    finally:
        await session.close()

if __name__ == "__main__":
    pytest.main(["-v", __file__])
